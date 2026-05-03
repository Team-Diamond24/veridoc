"""
LAYER 1: Document Ingestion & OCR
Converts all documents (PDF, DOCX, XLSX, images) into structured text chunks
with page references and confidence scores.
"""
import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional

logger = logging.getLogger(__name__)


def detect_file_type(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    type_map = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".xlsx": "xlsx",
        ".xls": "xlsx",
        ".jpg": "image",
        ".jpeg": "image",
        ".png": "image",
        ".tiff": "image",
        ".tif": "image",
        ".bmp": "image",
    }
    return type_map.get(ext, "unknown")


def extract_pdf(file_path: str) -> List[Dict[str, Any]]:
    """Extract text from PDF (typed or scanned). Returns list of chunk dicts."""
    chunks = []
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text and len(text.strip()) > 20:
                    # Typed PDF — high confidence
                    words = page.extract_words()
                    chunk_text = text.strip()
                    avg_confidence = 0.97  # Typed PDF

                    # Split into logical chunks by paragraph
                    paragraphs = [p.strip() for p in re.split(r'\n{2,}', chunk_text) if len(p.strip()) > 10]
                    if not paragraphs:
                        paragraphs = [chunk_text]

                    for i, para in enumerate(paragraphs):
                        bbox = None
                        if words:
                            # Estimate bbox from first word on page
                            try:
                                bbox = [
                                    float(words[0]["x0"]) if words else 0,
                                    float(words[0]["top"]) if words else 0,
                                    float(words[-1]["x1"]) if words else 600,
                                    float(words[-1]["bottom"]) if words else 800,
                                ]
                            except Exception:
                                bbox = [0, 0, 600, 800]

                        chunks.append({
                            "text": para,
                            "page": page_num,
                            "bbox": bbox or [0, 0, 600, 800],
                            "ocr_confidence": avg_confidence,
                            "source_label": f"Page {page_num}, Para {i+1}",
                            "chunk_index": len(chunks),
                        })
                else:
                    # Scanned page — try OCR
                    scanned_chunks = _ocr_page_image(file_path, page_num)
                    chunks.extend(scanned_chunks)

    except ImportError:
        logger.warning("pdfplumber not available, trying PyPDF2")
        chunks = _extract_pdf_pypdf2(file_path)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        chunks = _extract_pdf_pypdf2(file_path)

    return chunks if chunks else [{
        "text": "Document could not be fully parsed. Manual review required.",
        "page": 1,
        "bbox": [0, 0, 600, 800],
        "ocr_confidence": 0.3,
        "source_label": "Page 1",
        "chunk_index": 0,
    }]


def _extract_pdf_pypdf2(file_path: str) -> List[Dict[str, Any]]:
    chunks = []
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text and len(text.strip()) > 10:
                paragraphs = [p.strip() for p in re.split(r'\n{2,}', text) if len(p.strip()) > 10]
                for i, para in enumerate(paragraphs):
                    chunks.append({
                        "text": para,
                        "page": page_num,
                        "bbox": [0, 0, 600, 800],
                        "ocr_confidence": 0.92,
                        "source_label": f"Page {page_num}, Para {i+1}",
                        "chunk_index": len(chunks),
                    })
    except Exception as e:
        logger.error(f"PyPDF2 extraction error: {e}")
    return chunks


def _ocr_page_image(file_path: str, page_num: int) -> List[Dict[str, Any]]:
    """OCR a scanned PDF page using pytesseract."""
    chunks = []
    try:
        import pytesseract
        from PIL import Image
        import pdf2image

        images = pdf2image.convert_from_path(file_path, first_page=page_num, last_page=page_num)
        for img in images:
            data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT, lang='eng+hin')
            text_parts = []
            confidences = []
            for i, word in enumerate(data['text']):
                if word.strip() and int(data['conf'][i]) > 0:
                    text_parts.append(word)
                    confidences.append(int(data['conf'][i]) / 100)

            if text_parts:
                full_text = " ".join(text_parts)
                avg_conf = sum(confidences) / len(confidences) if confidences else 0.7
                paragraphs = [p.strip() for p in re.split(r'\n{2,}', full_text) if len(p.strip()) > 10]
                for i, para in enumerate(paragraphs or [full_text]):
                    chunks.append({
                        "text": para,
                        "page": page_num,
                        "bbox": [0, 0, 600, 800],
                        "ocr_confidence": round(avg_conf, 3),
                        "source_label": f"Page {page_num} (OCR), Para {i+1}",
                        "chunk_index": 0,
                    })
    except Exception as e:
        logger.warning(f"OCR failed for page {page_num}: {e}")
        chunks.append({
            "text": f"[Scanned page {page_num} — OCR unavailable. Manual review required.]",
            "page": page_num,
            "bbox": [0, 0, 600, 800],
            "ocr_confidence": 0.2,
            "source_label": f"Page {page_num} (Scanned)",
            "chunk_index": 0,
        })
    return chunks


def extract_docx(file_path: str) -> List[Dict[str, Any]]:
    """Extract text from Word documents."""
    chunks = []
    try:
        from docx import Document
        doc = Document(file_path)
        para_num = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if len(text) > 10:
                chunks.append({
                    "text": text,
                    "page": (para_num // 30) + 1,
                    "bbox": [0, 0, 600, 800],
                    "ocr_confidence": 0.99,
                    "source_label": f"Paragraph {para_num + 1}",
                    "chunk_index": para_num,
                })
                para_num += 1

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    chunks.append({
                        "text": row_text,
                        "page": (para_num // 30) + 1,
                        "bbox": [0, 0, 600, 800],
                        "ocr_confidence": 0.99,
                        "source_label": f"Table row {para_num + 1}",
                        "chunk_index": para_num,
                    })
                    para_num += 1
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")

    return chunks or [{
        "text": "DOCX document could not be parsed.",
        "page": 1,
        "bbox": [0, 0, 600, 800],
        "ocr_confidence": 0.3,
        "source_label": "Page 1",
        "chunk_index": 0,
    }]


def extract_xlsx(file_path: str) -> List[Dict[str, Any]]:
    """Extract data from Excel files as structured text."""
    chunks = []
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        chunk_idx = 0
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text = []
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(v) for v in row if v is not None]
                if row_vals:
                    rows_text.append(" | ".join(row_vals))
            if rows_text:
                chunks.append({
                    "text": f"[Sheet: {sheet_name}]\n" + "\n".join(rows_text[:50]),
                    "page": 1,
                    "bbox": [0, 0, 600, 800],
                    "ocr_confidence": 0.99,
                    "source_label": f"Sheet: {sheet_name}",
                    "chunk_index": chunk_idx,
                })
                chunk_idx += 1
    except Exception as e:
        logger.error(f"XLSX extraction error: {e}")
    return chunks


def extract_image(file_path: str) -> List[Dict[str, Any]]:
    """OCR an image file."""
    try:
        import pytesseract
        from PIL import Image

        img = Image.open(file_path)
        text = pytesseract.image_to_string(img, lang='eng+hin')
        conf_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        confidences = [int(c) for c in conf_data['conf'] if str(c).isdigit() and int(c) > 0]
        avg_conf = (sum(confidences) / len(confidences) / 100) if confidences else 0.7

        paragraphs = [p.strip() for p in re.split(r'\n{2,}', text) if len(p.strip()) > 10]
        return [
            {
                "text": para,
                "page": 1,
                "bbox": [0, 0, img.width, img.height],
                "ocr_confidence": round(avg_conf, 3),
                "source_label": f"Image OCR, Para {i+1}",
                "chunk_index": i,
            }
            for i, para in enumerate(paragraphs or [text])
        ]
    except Exception as e:
        logger.error(f"Image OCR error: {e}")
        return [{
            "text": "[Image — OCR unavailable. Manual review required.]",
            "page": 1,
            "bbox": [0, 0, 600, 800],
            "ocr_confidence": 0.2,
            "source_label": "Image",
            "chunk_index": 0,
        }]


def process_document(file_path: str) -> Dict[str, Any]:
    """
    Master function: detect type, extract text, return structured result.
    Returns:
        {
            file_type: str,
            page_count: int,
            avg_confidence: float,
            chunks: List[chunk_dict]
        }
    """
    file_type = detect_file_type(file_path)

    if file_type == "pdf":
        chunks = extract_pdf(file_path)
    elif file_type == "docx":
        chunks = extract_docx(file_path)
    elif file_type == "xlsx":
        chunks = extract_xlsx(file_path)
    elif file_type == "image":
        chunks = extract_image(file_path)
    else:
        chunks = [{
            "text": f"Unsupported file type: {Path(file_path).suffix}",
            "page": 1,
            "bbox": [0, 0, 600, 800],
            "ocr_confidence": 0.0,
            "source_label": "Unknown",
            "chunk_index": 0,
        }]

    # Re-index chunks
    for i, chunk in enumerate(chunks):
        chunk["chunk_index"] = i

    page_count = max((c["page"] for c in chunks), default=1)
    confidences = [c["ocr_confidence"] for c in chunks if c["ocr_confidence"] > 0]
    avg_confidence = round(sum(confidences) / len(confidences), 3) if confidences else 0.5

    return {
        "file_type": file_type,
        "page_count": page_count,
        "avg_confidence": avg_confidence,
        "chunks": chunks,
    }
